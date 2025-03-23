'''program to create an email as a .docx file.
   all paragraphs within the mail are new-line separated.
   the new-line of each paragrpah, a code is written in  "invisible ink",
   meaning that the color of the text is the same (white) as the background.
   the code contains bidding information for an undisclosed cartel deal 
   and is encoded as a vigenere cipher: encoding word: RANDOM
   
   R A N D O M R A N D O M R A N D O M R A N
   t h i s i s t h e c l e a r t e x t m s g

   encrypted via the following matrix:
    A  B C D E F ...
    -----------
    a| b c d e f ...
    b| c d e f g ...
    c| d e f g h ...
    d| e f g h i ...
    e| f g h i j ...
    ...
    ===> first line gives the letter of the encoding word;
         first line of the following row gives the letter in the code
         the encoding is the cross section between both
   '''

from docx import Document
from docx.shared import RGBColor
from invisible_ink.helpers import impexp as ie
from invisible_ink.helpers import vigenere as v

# =========================================================================
#                             USER INPUT:
# =========================================================================

CLEAR_TEXT_MSG = 'hello, my name is Phillip'
ENCODING_WORD = 'random'
MAIL_HEADER = 'Regarding the agreed-upon deal'
CIPH_CHARS_PER_PARAGRAPH = 10
DOC_NAME = 'invisile_ink.docx'

# =========================================================================
#                             FUNCTIONS:
# =========================================================================

def split_along_paragraphs(msg: str) -> list[str]:
    '''splits a txt in a list of strings, with paragraph changes (\n\n)
       as separator'''
    return msg.split(sep='\n\n')


def set_font_white(run) -> None:
    '''applies the color white to the current run'''
    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)


def set_font_as_before(run) -> None:
    '''set run color to the color inherited by the paragraph'''
    run.font.color.rgb = None


def generate_inv_ink_docx(
        mail_heading: str, mail_text: str, cipher: str, 
        ciph_chars_per_para: int
                          ) -> None:
    '''create a new docx file with an invisible ink cipher included'''
    inv_ink_docx = Document()
    inv_ink_docx.add_heading(mail_heading)
    mail_paragraphs = split_along_paragraphs(mail_text)
    for paragraph in mail_paragraphs:
        current_paragraph = inv_ink_docx.add_paragraph(paragraph)
        current_paragraph.add_run('\n')
        invisible_run = current_paragraph.add_run(cipher[:ciph_chars_per_para])
        set_font_white(invisible_run)
        cipher = cipher[ciph_chars_per_para:]
    return inv_ink_docx

# =========================================================================
#                             DRIVER CODE:
# =========================================================================

def main():
    '''main code'''
    encoder = v.VigenereEncoder(ENCODING_WORD, CLEAR_TEXT_MSG)
    encoded_msg = encoder.encode_vigenere()
    mail_msg = ie.import_txt('mail_text.txt')
    inv_ink_docx = generate_inv_ink_docx(
        MAIL_HEADER, mail_msg, encoded_msg, CIPH_CHARS_PER_PARAGRAPH
    )
    breakpoint()
    ie.export_dox('invisible_ink.docx', inv_ink_docx)

if __name__ == '__main__':
    main()
